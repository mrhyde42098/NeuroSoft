"""
app/infrastructure/export_service.py
=====================================
Exportadores DOCX y XLSX para el informe neuropsicológico.

Complementan `report_service.py` (PDF) con formatos editables:
- **DOCX**: informe completo editable en Word — mismo contenido que el PDF pero
  con estilos modificables por el profesional antes de imprimir/firmar.
- **XLSX**: matriz de puntajes — una hoja con los resultados (PD, PE, Z,
  interpretación) lista para análisis estadístico o anexo al informe.

Ambos usan `ReportData` como entrada, reutilizando la construcción que ya
hace `build_report_data_from_db`.
"""
from __future__ import annotations

import base64
import io
import logging
from datetime import date, datetime
from typing import Any

from app.infrastructure.report_service import ReportData

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════
# DOCX — Informe editable en Word
# ═══════════════════════════════════════════════════════════════

def generate_report_docx(data: ReportData) -> bytes:
    """
    Genera el informe completo en formato Word (.docx).
    Si `python-docx` no está instalado, levanta RuntimeError.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as e:
        raise RuntimeError(
            "python-docx no está instalado. Ejecuta: pip install python-docx"
        ) from e

    doc = Document()

    # ── Estilos base ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Márgenes A4 ajustados
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    _AZUL = RGBColor(0x1A, 0x56, 0x8C)
    _TEAL = RGBColor(0x2E, 0xC4, 0xB6)

    # ── Encabezado institucional ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.institucion_nombre or "Consultorio Neuropsicológico")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = _AZUL

    if any((data.institucion_nit, data.institucion_dir, data.institucion_tel)):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bits = []
        if data.institucion_nit:
            bits.append(f"NIT: {data.institucion_nit}")
        if data.institucion_dir:
            bits.append(data.institucion_dir)
        if data.institucion_tel:
            bits.append(f"Tel: {data.institucion_tel}")
        r = sub.add_run(" · ".join(bits))
        r.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("INFORME DE EVALUACIÓN NEUROPSICOLÓGICA")
    tr.bold = True
    tr.font.size = Pt(13)
    tr.font.color.rgb = _TEAL

    doc.add_paragraph()

    # ── Datos sociodemográficos ──
    _section_heading(doc, "1. DATOS SOCIODEMOGRÁFICOS", _AZUL)

    fecha_at = _fmt_date(data.fecha_atencion)
    fecha_nac = _fmt_date(data.fecha_nacimiento)
    pairs = [
        ("Nombre completo",     data.nombre_completo),
        ("Documento",           f"{data.tipo_documento} {data.numero_documento}"),
        ("Fecha de nacimiento", fecha_nac),
        ("Edad",                data.edad_display),
        ("Sexo",                data.sexo),
        ("Lateralidad",         data.lateralidad),
        ("Escolaridad",         data.escolaridad),
        ("Ocupación",           data.ocupacion),
        ("Ciudad",              data.ciudad),
        ("EPS",                 data.eps),
        ("Acompañante",         data.acompanante),
        ("Remitido por",        data.remite),
        ("Orden médica",        data.orden_no),
        ("Fecha de atención",   fecha_at),
        ("Protocolo aplicado",  data.protocolo),
    ]
    _key_value_table(doc, pairs)

    # ── Motivo de consulta ──
    if data.motivo_consulta:
        _section_heading(doc, "2. MOTIVO DE CONSULTA", _AZUL)
        doc.add_paragraph(data.motivo_consulta)

    # ── Antecedentes ──
    antec_pairs = [
        ("Patológicos / Médicos",    data.patologicos_medicos),
        ("Sensoriales / Motores",    data.sensoriales_motores),
        ("Psiquiátricos",            data.psiquiatricos),
        ("Farmacológicos",           data.farmacologicos),
        ("Traumáticos",              data.traumaticos),
        ("Quirúrgicos",              data.quirurgicos),
        ("Tóxicos",                  data.toxicos),
        ("Alérgicos",                data.alergicos),
        ("Terapéuticos previos",     data.terapeuticos),
        ("Paraclínicos",             data.paraclinicos),
        ("Familiares",               data.familiares),
    ]
    antec_nonempty = [(k, v) for k, v in antec_pairs if v]
    if antec_nonempty:
        _section_heading(doc, "3. ANTECEDENTES", _AZUL)
        _key_value_table(doc, antec_nonempty)

    # ── Funcionalidad / contexto ──
    func_pairs = [
        ("¿Con quién vive?",           data.vive_con),
        ("Actividades básicas",        data.abc),
        ("Contexto escolar / laboral", data.escolar_laboral),
        ("Patrón de sueño",            data.patron_sueno),
        ("Patrón de alimentación",     data.patron_alimentacion),
        ("Comportamiento y ánimo",     data.comportamiento_animo),
    ]
    func_nonempty = [(k, v) for k, v in func_pairs if v]
    if func_nonempty:
        _section_heading(doc, "4. FUNCIONALIDAD Y CONTEXTO", _AZUL)
        _key_value_table(doc, func_nonempty)

    # ── Observación clínica por dominios ──
    obs_pairs = [
        ("Apariencia y conducta",         data.obs_clinica_general),
        ("Atención y concentración",      data.obs_atencion),
        ("Memoria",                       data.obs_memoria),
        ("Praxias y gnosias",             data.obs_praxias_gnosias),
        ("Lenguaje",                      data.obs_lenguaje),
        ("Funciones ejecutivas",          data.obs_funciones_ejecutivas),
        ("Socioemocional",                data.obs_emociones),
        ("Cociente intelectual",          data.obs_ci),
        ("Funcionalidad global",          data.obs_funcionalidad),
    ]
    obs_nonempty = [(k, v) for k, v in obs_pairs if v]
    if obs_nonempty:
        _section_heading(doc, "5. OBSERVACIÓN CLÍNICA", _AZUL)
        for label, txt in obs_nonempty:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(txt)

    # ── Resultados cuantitativos ──
    if data.resultados:
        _section_heading(doc, "6. RESULTADOS CUANTITATIVOS", _AZUL)
        _results_table(doc, data.resultados)

    # ── Impresión diagnóstica ──
    if data.obs_impresion_dx:
        _section_heading(doc, "7. IMPRESIÓN DIAGNÓSTICA", _AZUL)
        doc.add_paragraph(data.obs_impresion_dx)

    # ── Recomendaciones ──
    if data.obs_recomendaciones:
        _section_heading(doc, "8. RECOMENDACIONES", _AZUL)
        doc.add_paragraph(data.obs_recomendaciones)

    # ── Aviso legal ──
    legal = doc.add_paragraph()
    lr = legal.add_run(data.aviso_legal)
    lr.italic = True
    lr.font.size = Pt(8.5)
    lr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Firma ──
    doc.add_paragraph()
    doc.add_paragraph()

    firma_p = doc.add_paragraph()
    firma_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Firma imagen si viene en base64
    if data.firma_base64:
        try:
            img_bytes = base64.b64decode(data.firma_base64.split(",")[-1])
            firma_p.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(5))
            doc.add_paragraph()
        except Exception as e:
            logger.warning("No se pudo insertar firma base64 en DOCX: %s", e)

    if data.profesional_nombre:
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = name_p.add_run(data.profesional_nombre)
        nr.bold = True

        if data.profesional_titulo:
            t_p = doc.add_paragraph()
            t_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            t_p.add_run(data.profesional_titulo).font.size = Pt(9)

        if data.profesional_registro:
            r_p = doc.add_paragraph()
            r_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = r_p.add_run(f"Registro Profesional: {data.profesional_registro}")
            rr.font.size = Pt(9)

    # Pie generado
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        f"Generado por NeuroSoft · {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Bytes ──
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ───────────── DOCX helpers ─────────────

def _section_heading(doc, text: str, color) -> None:
    """Título de sección con línea inferior tipo 'ruler'."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = color
    # Un párrafo de divisor sutil
    sep = doc.add_paragraph("─" * 60)
    for r in sep.runs:
        r.font.size = Pt(6)
        r.font.color.rgb = color


def _key_value_table(doc, pairs: list[tuple]) -> None:
    """Tabla key/value sin bordes, 2 columnas."""
    from docx.shared import Cm, Pt
    if not pairs:
        return
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.autofit = False
    for i, (k, v) in enumerate(pairs):
        cell_k = tbl.rows[i].cells[0]
        cell_v = tbl.rows[i].cells[1]
        cell_k.width = Cm(5.0)
        cell_v.width = Cm(11.0)
        # Key
        pk = cell_k.paragraphs[0]
        rk = pk.add_run(k)
        rk.bold = True
        rk.font.size = Pt(9.5)
        # Value
        pv = cell_v.paragraphs[0]
        rv = pv.add_run(str(v) if v is not None else "")
        rv.font.size = Pt(9.5)


def _results_table(doc, resultados: list[dict]) -> None:
    """Tabla de resultados cuantitativos con encabezados coloreados."""
    from docx.shared import Pt, RGBColor

    cols = ["Prueba", "Dominio", "PD", "PE", "Z", "Interpretación"]
    rows = len(resultados) + 1
    tbl = doc.add_table(rows=rows, cols=len(cols))
    tbl.style = "Light List Accent 1"

    # Cabecera
    for i, c in enumerate(cols):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].runs.clear() if False else None
        p = cell.paragraphs[0]
        r = p.add_run(c)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, res in enumerate(resultados, start=1):
        def _fmt(x):
            if x is None:
                return "—"
            try:
                return f"{float(x):.2f}"
            except Exception:
                return str(x)

        cells = tbl.rows[i].cells
        cells[0].text = str(res.get("test_nombre") or res.get("test_id") or "")
        cells[1].text = str(res.get("dominio_cognitivo") or "")
        cells[2].text = _fmt(res.get("puntaje_bruto"))
        cells[3].text = _fmt(res.get("puntaje_escalar"))
        cells[4].text = _fmt(res.get("z_equivalente"))
        cells[5].text = str(res.get("interpretacion") or "")
        for c in cells:
            for par in c.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)


# ═══════════════════════════════════════════════════════════════
# XLSX — Hoja de puntajes
# ═══════════════════════════════════════════════════════════════

def generate_evaluation_xlsx(data: ReportData) -> bytes:
    """
    Genera un XLSX con dos hojas:
      1. "Resumen"   — datos del paciente y protocolo
      2. "Puntajes"  — tabla completa de resultados (una fila por prueba)
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        raise RuntimeError(
            "openpyxl no está instalado. Ejecuta: pip install openpyxl"
        ) from e

    wb = Workbook()

    # Estilos compartidos
    header_fill = PatternFill("solid", fgColor="1A568C")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    label_font = Font(bold=True, size=10)
    value_font = Font(size=10)
    border_side = Side(style="thin", color="CCCCCC")
    thin_border = Border(left=border_side, right=border_side,
                         top=border_side, bottom=border_side)
    center = Alignment(horizontal="center", vertical="center")
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # ── Hoja 1: Resumen ──
    ws = wb.active
    ws.title = "Resumen"

    ws["A1"] = data.institucion_nombre or "Consultorio Neuropsicológico"
    ws["A1"].font = Font(bold=True, size=14, color="1A568C")
    ws.merge_cells("A1:B1")

    ws["A2"] = "INFORME DE EVALUACIÓN NEUROPSICOLÓGICA"
    ws["A2"].font = Font(bold=True, size=12, color="2EC4B6")
    ws.merge_cells("A2:B2")

    pairs = [
        ("Nombre completo",     data.nombre_completo),
        ("Documento",           f"{data.tipo_documento} {data.numero_documento}"),
        ("Fecha de nacimiento", _fmt_date(data.fecha_nacimiento)),
        ("Edad",                data.edad_display),
        ("Sexo",                data.sexo),
        ("Lateralidad",         data.lateralidad),
        ("Escolaridad",         data.escolaridad),
        ("Ocupación",           data.ocupacion),
        ("Ciudad",              data.ciudad),
        ("EPS",                 data.eps),
        ("Remitido por",        data.remite),
        ("Orden médica",        data.orden_no),
        ("Fecha de atención",   _fmt_date(data.fecha_atencion)),
        ("Protocolo",           data.protocolo),
        ("Motivo de consulta",  data.motivo_consulta),
        ("Impresión diagnóstica", data.obs_impresion_dx),
        ("Recomendaciones",     data.obs_recomendaciones),
        ("Profesional",         data.profesional_nombre),
        ("Registro profesional", data.profesional_registro),
    ]
    row = 4
    for k, v in pairs:
        c_key = ws.cell(row=row, column=1, value=k)
        c_val = ws.cell(row=row, column=2, value=v or "")
        c_key.font = label_font
        c_val.font = value_font
        c_val.alignment = left
        row += 1

    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 60

    # ── Hoja 2: Puntajes ──
    ws2 = wb.create_sheet(title="Puntajes")

    headers = [
        "Test ID", "Prueba", "Dominio", "PD (Puntaje Directo)",
        "PE (Escalar)", "Z equivalente", "Tipo métrica",
        "Interpretación", "Llave baremo",
    ]
    for j, h in enumerate(headers, start=1):
        cell = ws2.cell(row=1, column=j, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border

    resultados = data.resultados or []
    for i, res in enumerate(resultados, start=2):
        row_vals = [
            res.get("test_id"),
            res.get("test_nombre"),
            res.get("dominio_cognitivo"),
            _to_num(res.get("puntaje_bruto")),
            _to_num(res.get("puntaje_escalar")),
            _to_num(res.get("z_equivalente")),
            res.get("tipo_metrica"),
            res.get("interpretacion"),
            res.get("llave_baremo"),
        ]
        for j, val in enumerate(row_vals, start=1):
            cell = ws2.cell(row=i, column=j, value=val)
            cell.font = value_font
            cell.border = thin_border
            if j in (4, 5, 6):  # numéricos
                cell.alignment = center
                cell.number_format = "0.00"
            else:
                cell.alignment = left

    # Anchos
    col_widths = [14, 40, 22, 18, 14, 14, 16, 38, 18]
    for j, w in enumerate(col_widths, start=1):
        ws2.column_dimensions[get_column_letter(j)].width = w

    # Freeze header
    ws2.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════
# Helpers de formato
# ═══════════════════════════════════════════════════════════════

def _fmt_date(d: Any | None) -> str:
    if d is None:
        return ""
    if isinstance(d, (date, datetime)):
        return d.strftime("%d/%m/%Y")
    return str(d)


def _to_num(v: Any) -> float | None:
    if v is None or v == "":
        return None
    try:
        return float(v)
    except (TypeError, ValueError):
        return None
